import sys
import os
import io
import json
import asyncio
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from config import GITHUB_REPO, GITHUB_TOKEN, MCP_SUPPORTED_EXTENSIONS, MCP_MAX_FILE_BYTES
import httpx
from mcp.server.fastmcp import FastMCP

# Initialize the MCP Server
mcp = FastMCP("GitHub Knowledge Base Server")


def _get_runtime_repo() -> str:
    repo = os.getenv("GITHUB_REPO", "").strip() or GITHUB_REPO
    return repo.replace("https://github.com/", "").replace("http://github.com/", "").strip("/")


def _get_runtime_token() -> str:
    return os.getenv("GITHUB_TOKEN", "").strip() or GITHUB_TOKEN


# ─────────────────────────────────────────────────────────────
# File Type Parsers
# ─────────────────────────────────────────────────────────────

def _parse_pdf(raw_bytes: bytes) -> str:
    """Extract text from PDF binary content."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not available
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            return ""
    except Exception:
        return ""


def _parse_docx(raw_bytes: bytes) -> str:
    """Extract text from .docx binary content."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)
    except ImportError:
        return ""
    except Exception:
        return ""


def _parse_text(raw_bytes: bytes) -> str:
    """Parse plain text files with encoding fallback."""
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return ""


# Map extensions to parser types
_BINARY_EXTENSIONS = {".pdf", ".docx"}
_TEXT_EXTENSIONS = MCP_SUPPORTED_EXTENSIONS - _BINARY_EXTENSIONS


# ─────────────────────────────────────────────────────────────
# File Fetching
# ─────────────────────────────────────────────────────────────

async def fetch_file_content(client: httpx.AsyncClient, file_path: str, file_size: int) -> dict:
    github_repo = _get_runtime_repo()
    github_token = _get_runtime_token()

    """
    Fetch and parse content for a single file from GitHub.
    Handles both text and binary (PDF/docx) formats.
    """
    # Skip files that are too large
    if file_size > MCP_MAX_FILE_BYTES:
        return {"content": "", "source": file_path, "technology": "General", "doc_name": ""}

    suffix = Path(file_path).suffix.lower()
    is_binary = suffix in _BINARY_EXTENSIONS

    # Build raw URL
    raw_url = f"https://raw.githubusercontent.com/{github_repo}/main/{file_path}"
    headers = {}
    if github_token:
        headers["Authorization"] = f"token {github_token}"

    response = await client.get(raw_url, headers=headers)

    # Fallback to master branch
    if response.status_code == 404:
        raw_url_master = f"https://raw.githubusercontent.com/{github_repo}/master/{file_path}"
        response = await client.get(raw_url_master, headers=headers)

    response.raise_for_status()

    # Parse based on file type
    if suffix == ".pdf":
        content = _parse_pdf(response.content)
    elif suffix == ".docx":
        content = _parse_docx(response.content)
    else:
        content = _parse_text(response.content)

    content = content.strip()

    # Extract metadata from path
    parts = Path(file_path).parts
    tech_category = parts[-2] if len(parts) >= 2 else "General"
    doc_name = Path(file_path).stem

    return {
        "content": content,
        "source": file_path,
        "technology": tech_category,
        "doc_name": doc_name,
        "file_type": suffix,
    }


# ─────────────────────────────────────────────────────────────
# MCP Tool
# ─────────────────────────────────────────────────────────────

@mcp.tool()
async def fetch_repo_contents() -> str:
    """
    Fetch all supported documents from the configured GitHub repository.
    Supports: .md, .txt, .log, .sql, .pdf, .docx, .csv, .json, .xml, .yaml, .py, .sh, .conf, .cfg, .ini
    Returns a JSON list of dictionaries containing file content and metadata.
    """
    github_repo = _get_runtime_repo()
    github_token = _get_runtime_token()
    if not github_repo:
        raise ValueError("GITHUB_REPO environment variable is not set. Please set it to 'owner/repo'.")

    api_url = f"https://api.github.com/repos/{github_repo}/git/trees/main?recursive=1"

    headers = {
        "Accept": "application/vnd.github.v3+json",
        "User-Agent": "DataGuru-MCP-Server"
    }
    if github_token:
        headers["Authorization"] = f"token {github_token}"

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.get(api_url, headers=headers)

        # Fallback to master if main branch is not found
        if response.status_code == 404:
            api_url_master = f"https://api.github.com/repos/{github_repo}/git/trees/master?recursive=1"
            response = await client.get(api_url_master, headers=headers)

        response.raise_for_status()
        tree_data = response.json().get("tree", [])

        # Filter for supported file types
        supported_files = []
        for item in tree_data:
            if item["type"] != "blob":
                continue
            suffix = Path(item["path"]).suffix.lower()
            if suffix in MCP_SUPPORTED_EXTENSIONS:
                file_size = item.get("size", 0)
                supported_files.append((item["path"], file_size))

        print(f"Found {len(supported_files)} supported files in {github_repo}.", file=sys.stderr)

        # Download files concurrently (batch to avoid overwhelming the API)
        batch_size = 10
        documents = []
        for i in range(0, len(supported_files), batch_size):
            batch = supported_files[i:i + batch_size]
            tasks = [fetch_file_content(client, path, size) for path, size in batch]
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            documents.extend(batch_results)

        # Filter out failures and empty documents
        successful_docs = [
            doc for doc in documents
            if not isinstance(doc, Exception) and isinstance(doc, dict) and doc.get("content")
        ]

        print(f"Successfully fetched {len(successful_docs)} documents from {github_repo} via MCP.", file=sys.stderr)
        return json.dumps(successful_docs)


if __name__ == "__main__":
    print("Starting GitHub Knowledge Base MCP Server...", file=sys.stderr)
    mcp.run()
